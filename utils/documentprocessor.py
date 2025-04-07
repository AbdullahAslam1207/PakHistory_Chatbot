import fitz  # PyMuPDF for PDF processing
import docx  # python-docx for Word files
from langchain_community.document_loaders import TextLoader
from langchain.schema import Document

class DocumentProcessor:
    def __init__(self):
        self.documents = []
        self.text_files=[]
        self.pdf_files=[]
        self.docx_files=[]

    def add_files(self,files):
        for link in files:
            if link.endswith(".pdf"):
                self.pdf_files.append(link)
            elif link.endswith(".txt"):
                self.text_files.append(link)
            elif link.endswith(".docx"):
                self.docx_files.append(link)


    def process_pdfs(self):
        """Processes PDF files and appends their text content to the documents list."""
        for pdf in self.pdf_files:
            doc = fitz.open(pdf)
            text = ""
            for page in doc:
                text += page.get_text()
            self.documents.append(Document(page_content=text))
            doc.close()

    def process_text_files(self):
        """Processes text files and appends their content to the documents list."""
        for text_file in self.text_files:
            try:
                with open(text_file, 'r', encoding='utf-8') as file:
                    text = file.read()  # Read the entire content of the file
                self.documents.append(Document(page_content=text))
            except Exception as e:
                print(f"Error processing file {text_file}: {e}")


    def process_word_files(self):
        """Processes Word (.docx) files and appends their content to the documents list."""
        for docx_file in self.docx_files:
            doc = docx.Document(docx_file)
            self.text = "\n".join([para.text for para in doc.paragraphs])  # Reset and store content
            self.documents.append(Document(page_content=self.text))

    def get_documents(self):
        """Returns the processed documents."""
        return self.documents
    
